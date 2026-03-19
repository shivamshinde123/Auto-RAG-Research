"""Google Drive data source connector.

Reads PDF and DOCX files from a configured Google Drive folder
using the Drive API v3 with OAuth2 credentials.
"""

import io
import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document

from src.data_sources import register
from src.data_sources.base import BaseDataSource

logger = logging.getLogger(__name__)

SUPPORTED_MIMES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


@register("gdrive")
class GdriveDataSource(BaseDataSource):

    def validate_config(self) -> bool:
        folder_id = self.config.get("folder_id")
        if not folder_id:
            raise ValueError("gdrive config missing required 'folder_id' field")
        creds_path = self.config.get("credentials_path", ".secrets/credentials.json")
        if not Path(creds_path).exists():
            raise ValueError(
                f"Google Drive credentials not found at '{creds_path}'. "
                "Place your OAuth2 credentials.json in .secrets/"
            )
        return True

    def health_check(self) -> bool:
        self.validate_config()
        service = self._build_service()
        folder_id = self.config["folder_id"]
        try:
            service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                pageSize=1,
                fields="files(id)",
            ).execute()
        except Exception as e:
            raise ConnectionError(f"Cannot access Google Drive folder '{folder_id}': {e}")
        logger.info("health_check passed: Google Drive folder %s accessible", folder_id)
        return True

    def load(self) -> List[Document]:
        self.validate_config()
        service = self._build_service()
        folder_id = self.config["folder_id"]
        documents: List[Document] = []

        # List files in folder (with pagination)
        query = f"'{folder_id}' in parents and trashed=false"
        files = []
        page_token = None
        while True:
            kwargs = dict(
                q=query,
                pageSize=100,
                fields="nextPageToken, files(id, name, mimeType, modifiedTime)",
            )
            if page_token:
                kwargs["pageToken"] = page_token
            results = service.files().list(**kwargs).execute()
            files.extend(results.get("files", []))
            page_token = results.get("nextPageToken")
            if not page_token:
                break

        for file_info in files:
            mime = file_info.get("mimeType", "")
            if mime not in SUPPORTED_MIMES:
                logger.debug("Skipping unsupported file type: %s (%s)", file_info["name"], mime)
                continue

            try:
                content = service.files().get_media(fileId=file_info["id"]).execute()
                file_type = SUPPORTED_MIMES[mime]

                if file_type == "pdf":
                    docs = self._extract_pdf(content, file_info)
                elif file_type == "docx":
                    docs = self._extract_docx(content, file_info)
                elif file_type == "txt":
                    docs = self._extract_txt(content, file_info)
                else:
                    continue

                documents.extend(docs)
            except Exception as e:
                logger.error("Failed to process %s: %s", file_info["name"], e)

        logger.info("Loaded %d documents from Google Drive folder %s", len(documents), folder_id)
        return documents

    def _build_service(self):
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build

        creds_path = self.config.get("credentials_path", ".secrets/credentials.json")
        token_path = self.config.get("token_path", ".secrets/token.json")

        creds = None
        if Path(token_path).exists():
            creds = Credentials.from_authorized_user_file(
                token_path, ["https://www.googleapis.com/auth/drive.readonly"]
            )

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                from google.auth.transport.requests import Request
                creds.refresh(Request())
            else:
                from google_auth_oauthlib.flow import InstalledAppFlow
                flow = InstalledAppFlow.from_client_secrets_file(
                    creds_path, ["https://www.googleapis.com/auth/drive.readonly"]
                )
                creds = flow.run_local_server(port=0)
            Path(token_path).parent.mkdir(parents=True, exist_ok=True)
            Path(token_path).write_text(creds.to_json())

        return build("drive", "v3", credentials=creds)

    def _extract_pdf(self, content: bytes, file_info: dict) -> List[Document]:
        import fitz

        docs: List[Document] = []
        doc = fitz.open(stream=content, filetype="pdf")
        try:
            for page_num in range(len(doc)):
                text = doc[page_num].get_text()
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={
                            "source": f"gdrive://{file_info['id']}",
                            "source_type": "gdrive",
                            "file_name": file_info["name"],
                            "drive_file_id": file_info["id"],
                            "last_modified": file_info.get("modifiedTime", ""),
                            "page_number": page_num + 1,
                        },
                    ))
        finally:
            doc.close()
        return docs

    def _extract_txt(self, content: bytes, file_info: dict) -> List[Document]:
        text = content.decode("utf-8", errors="replace")
        if not text.strip():
            return []
        return [Document(
            page_content=text,
            metadata={
                "source": f"gdrive://{file_info['id']}",
                "source_type": "gdrive",
                "file_name": file_info["name"],
                "drive_file_id": file_info["id"],
                "last_modified": file_info.get("modifiedTime", ""),
            },
        )]

    def _extract_docx(self, content: bytes, file_info: dict) -> List[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text.strip():
            return []
        return [Document(
            page_content=text,
            metadata={
                "source": f"gdrive://{file_info['id']}",
                "source_type": "gdrive",
                "file_name": file_info["name"],
                "drive_file_id": file_info["id"],
                "last_modified": file_info.get("modifiedTime", ""),
            },
        )]
